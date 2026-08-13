"""
redact_pii.py — Enterprise PII Redaction Engine
=================================================
A high-precision PII redaction tool that processes .docx files using
Microsoft Presidio + spaCy for entity detection, Faker for consistent
replacement generation, and python-docx for format-preserving output.

Architecture:
    1. AnalyzerEngine (Presidio + spaCy) detects PII spans in text.
    2. A 1:1 ConsistentMapper ensures identical PII values always map
       to the same fake replacement across the entire document.
    3. The DocxProcessor walks every paragraph, header, footer, and
       table cell — replacing text at the *run* level to preserve
       all original formatting (bold, italic, font size, color, etc.).
    4. An EvaluationEngine computes Precision, Recall, and Accuracy
       against a user-supplied ground-truth annotation set.

Usage:
    python redact_pii.py --input <file.docx> --output <redacted.docx>
    python redact_pii.py --input <file.docx> --output <redacted.docx> \
                         --ground-truth <annotations.json> --evaluate
"""

# ──────────────────────────────────────────────────────────────────────
# Imports
# ──────────────────────────────────────────────────────────────────────
import argparse
import json
import re
import sys
import copy
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

from docx import Document                       # python-docx
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn                     # XML namespace helper
from docx.oxml import OxmlElement
from faker import Faker                         # Realistic fake data
from presidio_analyzer import (                 # Microsoft Presidio
    AnalyzerEngine,
    PatternRecognizer,
    Pattern,
    RecognizerResult,
)
from presidio_analyzer.nlp_engine import NlpEngineProvider
from PIL import Image
import pytesseract

# ──────────────────────────────────────────────────────────────────────
# Constants & Configuration
# ──────────────────────────────────────────────────────────────────────

# All PII entity types the engine is configured to detect.
# To add a new type: (1) add it here, (2) register a recognizer in
# _build_analyzer(), (3) add a faker generator in ConsistentMapper.
SUPPORTED_PII_TYPES: List[str] = [
    "PERSON",              # Full names
    "EMAIL_ADDRESS",       # Email addresses
    "PHONE_NUMBER",        # Phone numbers (US / international)
    "ORGANIZATION",        # Company / organisation names
    "ADDRESS",             # Physical / mailing addresses
    "US_SSN",              # Social Security Numbers
    "CREDIT_CARD",         # Credit / debit card numbers
    "DATE_OF_BIRTH",       # Dates of birth
    "IP_ADDRESS",          # IPv4 / IPv6 addresses
]

# Minimum confidence score (0‒1) a Presidio result must reach to be
# treated as a true detection.  Raising this reduces false positives
# but may increase false negatives.
CONFIDENCE_THRESHOLD: float = 0.45

# Per-entity threshold overrides.  Some entity types (e.g. phone
# numbers in table cells with no surrounding context) produce lower
# confidence scores, so we allow a lower floor for those.
PER_ENTITY_THRESHOLDS: Dict[str, float] = {
    "PHONE_NUMBER": 0.35,
    "US_SSN":       0.35,
    "CREDIT_CARD":  0.35,
    "IP_ADDRESS":   0.35,
    "EMAIL_ADDRESS": 0.35,
}

# Deny-list: terms that commonly trigger false positives from spaCy NER.
# These are case-insensitive — any detected PII whose text matches an
# entry here (after normalising whitespace & case) will be suppressed.
DENY_LIST: Set[str] = {
    # Generic labels that spaCy misclassifies as PERSON
    "email", "email address", "phone", "phone number", "full name",
    "home address", "address", "contact name", "employer",
    "summary", "appendix", "appendix a", "appendix b",
    "field", "details", "notes", "date", "order",
    # Generic terms that spaCy misclassifies as ORGANIZATION
    "social security", "social security number", "engineering",
    "ip", "it", "hr", "qa", "the", "service", "service agreement",
    "confidential", "personal information", "employment details",
    "emergency contact",
    # Financial / legal terms common in prospectus documents
    "section", "chapter", "part", "clause", "article", "schedule",
    "regulation", "regulations", "act", "rule", "rules",
    "offer", "issue", "bid", "allotment", "subscription",
    "equity", "share", "shares", "capital", "stock",
    "board", "committee", "directors", "management",
    "promoter", "promoters", "group", "investor", "investors",
    "company", "entity", "firm", "trust", "fund",
    "fiscal", "financial", "quarter", "annual", "period",
    "net", "gross", "total", "revenue", "profit", "loss",
    "march", "april", "may", "june", "july", "august",
    "september", "october", "november", "december",
    "january", "february",
    "red herring", "prospectus", "red herring prospectus",
    "book built", "anchor", "retail", "non-institutional",
    "general", "specific", "material", "aggregate",
    # Indian geographic / legal / industry terms
    "taluka", "taluka-khed", "kanjurmarg", "erandawane",
    "marg", "lok sabha", "rajya sabha", "corrigenda",
    "challan", "distriparks", "unpai", "deen dayal",
    "urja suraksha", "cap price", "continuous transposed conductors",
    "gopal bo", "appasaheb marathe marg",
}

# Minimum digit count for PHONE_NUMBER detections.  Strings with
# fewer digits than this are suppressed (catches reference numbers
# like "16949" that Presidio flags at low confidence).
MIN_PHONE_DIGITS: int = 7

# Seed for Faker — set to a fixed int for reproducible fake data.
FAKER_SEED: Optional[int] = 42


# ──────────────────────────────────────────────────────────────────────
# 1. Analyzer Builder  —  Presidio + spaCy + custom recognizers
# ──────────────────────────────────────────────────────────────────────

def _build_analyzer() -> AnalyzerEngine:
    """
    Construct a Presidio AnalyzerEngine backed by the spaCy NLP model
    and augmented with custom pattern-based recognizers for entity
    types that Presidio does not cover out-of-the-box.

    Returns:
        A fully configured AnalyzerEngine instance.
    """

    # --- spaCy NLP backend ------------------------------------------------
    # We use spaCy's `en_core_web_lg` model for higher accuracy on named
    # entities (PERSON, ORG, GPE).  Falls back to `en_core_web_sm` if lg
    # is not installed, but accuracy will degrade.
    nlp_configuration = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}],
    }

    try:
        nlp_engine = NlpEngineProvider(
            nlp_configuration=nlp_configuration
        ).create_engine()
    except OSError:
        # Fallback to small model
        nlp_configuration["models"][0]["model_name"] = "en_core_web_sm"
        nlp_engine = NlpEngineProvider(
            nlp_configuration=nlp_configuration
        ).create_engine()

    analyzer = AnalyzerEngine(
        nlp_engine=nlp_engine,
        supported_languages=["en"],
    )

    # --- Custom Recognizer: DATE_OF_BIRTH --------------------------------
    # Presidio has no built-in DOB recognizer. We define common date
    # patterns that appear in enterprise documents.
    dob_patterns = [
        Pattern(
            "DOB_MDY_SLASH",
            r"\b(0[1-9]|1[0-2])/(0[1-9]|[12]\d|3[01])/(19|20)\d{2}\b",
            0.70,
        ),
        Pattern(
            "DOB_DMY_SLASH",
            r"\b(0[1-9]|[12]\d|3[01])/(0[1-9]|1[0-2])/(19|20)\d{2}\b",
            0.60,
        ),
        Pattern(
            "DOB_WRITTEN",
            r"\b(?:January|February|March|April|May|June|July|August|"
            r"September|October|November|December)\s+"
            r"(0?[1-9]|[12]\d|3[01]),?\s+(19|20)\d{2}\b",
            0.85,
        ),
        Pattern(
            "DOB_ISO",
            r"\b(19|20)\d{2}-(0[1-9]|1[0-2])-(0[1-9]|[12]\d|3[01])\b",
            0.75,
        ),
        Pattern(
            "DOB_DMY_DASH",
            r"\b(0[1-9]|[12]\d|3[01])-(0[1-9]|1[0-2])-(19|20)\d{2}\b",
            0.65,
        ),
    ]
    dob_recognizer = PatternRecognizer(
        supported_entity="DATE_OF_BIRTH",
        patterns=dob_patterns,
        supported_language="en",
        name="DateOfBirthRecognizer",
        context=["born", "birth", "dob", "date of birth", "birthday"],
    )
    analyzer.registry.add_recognizer(dob_recognizer)

    # --- Custom Recognizer: US_SSN (supplement Presidio built-in) ---------
    # Presidio's built-in US_SSN validator rejects area numbers that were
    # never officially assigned (e.g. 123-xx-xxxx).  In enterprise docs
    # ANY 3-2-4 digit pattern should be treated as sensitive, so we add a
    # broader pattern that skips the area-code validation.
    ssn_patterns = [
        Pattern(
            "SSN_BROAD",
            r"\b\d{3}-\d{2}-\d{4}\b",
            0.60,
        ),
        Pattern(
            "SSN_NO_DASHES",
            r"\b\d{9}\b",
            0.30,
        ),
    ]
    ssn_recognizer = PatternRecognizer(
        supported_entity="US_SSN",
        patterns=ssn_patterns,
        supported_language="en",
        name="BroadSSNRecognizer",
        context=["ssn", "social security", "social security number",
                 "social", "tax id", "taxpayer"],
    )
    analyzer.registry.add_recognizer(ssn_recognizer)

    # --- Custom Recognizer: ADDRESS (supplement spaCy GPE) ----------------
    # spaCy detects GPE (geo-political entities) but misses full mailing
    # addresses.  This pattern catches street addresses.
    # CRITICAL: The street-type suffix (Street, Road, etc.) is MANDATORY
    # to avoid matching patterns like "10 customers contributed to".
    address_patterns = [
        Pattern(
            "US_STREET_ADDRESS",
            r"\b\d{1,5}\s+"
            r"(?:[A-Z][a-z]+\s+){1,3}"
            r"(?:Street|St|Avenue|Ave|Boulevard|Blvd|Drive|Dr|"
            r"Road|Rd|Lane|Ln|Way|Court|Ct|Place|Pl|Circle|Cir)"
            r"\.?"
            r"(?:\s*,?\s*(?:Suite|Ste|Apt|Unit|#)\s*\d+)?"
            r"(?:\s*,?\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)?"
            r"(?:\s*,?\s*[A-Z]{2})?"
            r"(?:\s+\d{5}(?:-\d{4})?)?\b",
            0.70,
        ),
        # Indian-style addresses (Village/Taluka/District patterns)
        Pattern(
            "IN_ADDRESS",
            r"\b\d{1,5}(?:/\d{1,5})*\s*,?\s*"
            r"(?:Village|Taluka|District|Sector|Plot|Survey|Block)"
            r"\s+[A-Z][A-Za-z\s,\-]+"
            r"(?:\d{3}\s*\d{3})?\b",
            0.70,
        ),
    ]
    address_recognizer = PatternRecognizer(
        supported_entity="ADDRESS",
        patterns=address_patterns,
        supported_language="en",
        name="AddressRecognizer",
        context=["address", "located", "residing", "mailing", "street"],
    )
    analyzer.registry.add_recognizer(address_recognizer)

    # --- Custom Recognizer: ORGANIZATION (boost spaCy ORG) ----------------
    # spaCy's NER already detects ORG, but we add suffix patterns to
    # catch company names that spaCy occasionally misses.
    org_patterns = [
        Pattern(
            "ORG_SUFFIX",
            r"\b[A-Z][A-Za-z&\-\s]{1,50}\s+"
            r"(?:Inc|LLC|Ltd|Corp|Corporation|Co|Company|Group|"
            r"Partners|Holdings|Enterprises|Solutions|Technologies|"
            r"Consulting|Services|International|Associates|GmbH|"
            r"Pvt\.?\s*Ltd\.?|LLP|PLC|AG|SA|NV|BV)"
            r"\.?\b",
            0.75,
        ),
    ]
    org_recognizer = PatternRecognizer(
        supported_entity="ORGANIZATION",
        patterns=org_patterns,
        supported_language="en",
        name="OrganizationPatternRecognizer",
        context=["company", "employer", "firm", "organization", "client"],
    )
    analyzer.registry.add_recognizer(org_recognizer)

    return analyzer


# ──────────────────────────────────────────────────────────────────────
# 2. Consistent 1:1 Mapper  —  Faker-based fake data generator
# ──────────────────────────────────────────────────────────────────────

class ConsistentMapper:
    """
    Maintains a deterministic 1:1 mapping from each unique PII value
    to a realistic Faker-generated replacement.

    If "Rashi Patil" is detected as PERSON in paragraph 1, the mapper
    returns (say) "Margaret Johnson" and will return the *same* string
    for every subsequent occurrence of "Rashi Patil" in the document.

    Thread Safety:
        This class is NOT thread-safe.  Instantiate one per document.
    """

    def __init__(self, seed: Optional[int] = FAKER_SEED) -> None:
        self._faker = Faker()
        if seed is not None:
            Faker.seed(seed)
        # dict[entity_type][original_value] → replacement_value
        self._map: Dict[str, Dict[str, str]] = {}

    # ---- public API -----------------------------------------------------

    def get_replacement(self, entity_type: str, original: str) -> str:
        """
        Return the fake replacement for *original*.  If *original* has
        been seen before under the same *entity_type*, return the cached
        value; otherwise generate a new one and cache it.
        """
        type_map = self._map.setdefault(entity_type, {})
        # Normalise whitespace for matching consistency
        key = " ".join(original.split())
        if key not in type_map:
            type_map[key] = self._generate(entity_type)
        return type_map[key]

    def get_full_map(self) -> Dict[str, Dict[str, str]]:
        """Return a copy of the entire mapping dictionary (for logging)."""
        return {k: dict(v) for k, v in self._map.items()}

    # ---- private generators ---------------------------------------------

    def _generate(self, entity_type: str) -> str:
        """
        Dispatch to the appropriate Faker provider based on entity type.

        To add a NEW PII type:
            1. Add a branch here.
            2. Register the entity in SUPPORTED_PII_TYPES.
            3. Add a Presidio recognizer in _build_analyzer().
        """
        generators = {
            "PERSON":         self._gen_person,
            "EMAIL_ADDRESS":  self._gen_email,
            "PHONE_NUMBER":   self._gen_phone,
            "ORGANIZATION":   self._gen_org,
            "ADDRESS":        self._gen_address,
            "US_SSN":         self._gen_ssn,
            "CREDIT_CARD":    self._gen_credit_card,
            "DATE_OF_BIRTH":  self._gen_dob,
            "IP_ADDRESS":     self._gen_ip,
        }
        gen_fn = generators.get(entity_type, self._gen_generic)
        return gen_fn()

    def _gen_person(self) -> str:
        return self._faker.name()

    def _gen_email(self) -> str:
        return self._faker.email()

    def _gen_phone(self) -> str:
        return self._faker.phone_number()

    def _gen_org(self) -> str:
        return self._faker.company()

    def _gen_address(self) -> str:
        # Single-line address for inline replacement
        return self._faker.address().replace("\n", ", ")

    def _gen_ssn(self) -> str:
        return self._faker.ssn()

    def _gen_credit_card(self) -> str:
        return self._faker.credit_card_number()

    def _gen_dob(self) -> str:
        return self._faker.date_of_birth(
            minimum_age=18, maximum_age=90
        ).strftime("%m/%d/%Y")

    def _gen_ip(self) -> str:
        return self._faker.ipv4()

    def _gen_generic(self) -> str:
        """Fallback for any unrecognised entity type."""
        return f"[REDACTED-{self._faker.uuid4()[:8]}]"


# ──────────────────────────────────────────────────────────────────────
# 3. Redaction Engine  —  ties Presidio analysis to the mapper
# ──────────────────────────────────────────────────────────────────────

@dataclass
class RedactionResult:
    """Container for a single redaction operation on a text string."""
    original_text: str
    redacted_text: str
    detections: List[RecognizerResult] = field(default_factory=list)


class RedactionEngine:
    """
    Orchestrator that:
      1. Passes raw text through the Presidio AnalyzerEngine.
      2. Filters results by confidence threshold.
      3. Resolves overlapping spans (keeps highest-confidence).
      4. Replaces detected PII using the ConsistentMapper.
    """

    def __init__(
        self,
        analyzer: Optional[AnalyzerEngine] = None,
        mapper: Optional[ConsistentMapper] = None,
        confidence_threshold: float = CONFIDENCE_THRESHOLD,
    ) -> None:
        self.analyzer = analyzer or _build_analyzer()
        self.mapper = mapper or ConsistentMapper()
        self.threshold = confidence_threshold

    def redact(self, text: str) -> RedactionResult:
        """
        Analyse *text* for PII and return a RedactionResult with the
        redacted string and the list of detections used.
        """
        if not text or not text.strip():
            return RedactionResult(
                original_text=text, redacted_text=text
            )

        # Step 1: Run Presidio analysis
        raw_results: List[RecognizerResult] = self.analyzer.analyze(
            text=text,
            entities=SUPPORTED_PII_TYPES,
            language="en",
        )

        # Step 2: Filter by confidence threshold (with per-entity overrides)
        filtered = [
            r for r in raw_results
            if r.score >= PER_ENTITY_THRESHOLDS.get(
                r.entity_type, self.threshold
            )
        ]

        # Step 2b: Remove deny-listed false positives
        filtered = [
            r for r in filtered
            if " ".join(text[r.start : r.end].split()).lower()
            not in DENY_LIST
        ]

        # Step 2c: PERSON heuristic filter
        # Suppress single-word detections that are likely false positives:
        #   - All-uppercase (abbreviations like BESS, EMIS, SCRR)
        #   - Fewer than 4 characters (Li, GW, A.Y.)
        #   - Contains digits or special chars (not a real name)
        #   - Known geographic/legal terms (Taluka, Branch, etc.)
        _GEO_LEGAL_TERMS = {
            "taluka", "branch", "account", "apartment", "schedule",
            "circular", "slip", "bill", "grill", "margin", "conductor",
            "participant", "facility", "park", "industrial", "sector",
            "village", "district", "division", "unit", "block",
            "marg", "nagar", "wadi", "pura", "abad", "ganj",
            "chowk", "chakan", "cantonment", "sabha",
            "conductors", "transposed", "continuous",
            "suraksha", "urja", "dayal", "deen",
            "distriparks", "corrigenda", "challan",
        }
        def _is_likely_person(r: RecognizerResult) -> bool:
            if r.entity_type != "PERSON":
                return True  # not a PERSON, don't filter
            val = text[r.start : r.end].strip()
            # Reject all-uppercase (abbreviations)
            if val.isupper() and len(val) > 1:
                return False
            # Reject very short tokens
            if len(val) <= 2:
                return False
            # Reject if contains digits
            if any(c.isdigit() for c in val):
                return False
            # Reject known geographic/legal terms in the value
            val_words = val.lower().split()
            if any(w in _GEO_LEGAL_TERMS for w in val_words):
                return False
            # Reject abbreviation patterns like "A.Y.", "W&C"
            if re.match(r'^[A-Z][.&/][A-Z]?[.]?$', val):
                return False
            # Reject single-word entries that aren't plausible names
            # (Plausible = title-case, ≥ 4 chars)
            if len(val_words) == 1 and len(val) < 4:
                return False
            return True

        filtered = [r for r in filtered if _is_likely_person(r)]

        # Step 2d: Phone-number minimum-digit filter
        # Real phone numbers have 10+ digits, or 7+ digits with a
        # country-code prefix (+).  This filters out reference numbers.
        def _is_likely_phone(r: RecognizerResult) -> bool:
            if r.entity_type != "PHONE_NUMBER":
                return True
            val = text[r.start : r.end]
            digit_count = sum(c.isdigit() for c in val)
            has_plus = val.strip().startswith("+")
            has_parens = "(" in val
            # Reject date-like reference numbers (YYYYMMDD-NN)
            if re.match(r'^\d{8}-\d+$', val.strip()):
                return False
            # Accept if 10+ digits, or 7+ with country code/formatting
            if digit_count >= 10:
                return True
            if digit_count >= 7 and (has_plus or has_parens):
                return True
            return False

        filtered = [r for r in filtered if _is_likely_phone(r)]

        # Step 2e: Anti-cascade — skip values that are already
        # Faker-generated replacements from previous redactions.
        # This prevents replacement phone/email values from being
        # re-detected and re-mapped in later paragraphs.
        existing_replacements = set()
        for type_map in self.mapper.get_full_map().values():
            existing_replacements.update(type_map.values())
        filtered = [
            r for r in filtered
            if text[r.start : r.end] not in existing_replacements
        ]

        # Step 2f: Context Enhancement — reduce false positives for generic numbers
        # Ignore numbers if preceded by generic reference terms (e.g., Order, Ticket, Invoice)
        IGNORE_CONTEXTS = {"order", "ticket", "invoice", "id", "ref"}
        def _is_valid_context(r: RecognizerResult) -> bool:
            # Check trailing 25 characters for ignore contexts
            context_window = text[max(0, r.start - 25) : r.start].lower()
            if any(ctx in context_window for ctx in IGNORE_CONTEXTS):
                return False
            return True

        filtered = [r for r in filtered if _is_valid_context(r)]

        # Step 3: Resolve overlapping detections — keep highest score
        resolved = self._resolve_overlaps(filtered)

        # Step 4: Build redacted text by replacing spans right-to-left
        #         (right-to-left avoids offset shifts).
        redacted = text
        for detection in sorted(resolved, key=lambda d: d.start, reverse=True):
            original_value = text[detection.start : detection.end]

            # Handle possessive forms: "Rashi Patil's" should map
            # consistently with "Rashi Patil" and re-append the suffix.
            suffix = ""
            clean_value = original_value
            if detection.entity_type == "PERSON":
                for poss in ("'s", "\u2019s", "'s"):
                    if clean_value.endswith(poss):
                        suffix = poss
                        clean_value = clean_value[: -len(poss)]
                        break

            replacement = self.mapper.get_replacement(
                detection.entity_type, clean_value
            ) + suffix

            redacted = (
                redacted[: detection.start]
                + replacement
                + redacted[detection.end :]
            )

        return RedactionResult(
            original_text=text,
            redacted_text=redacted,
            detections=resolved,
        )

    @staticmethod
    def _resolve_overlaps(
        results: List[RecognizerResult],
    ) -> List[RecognizerResult]:
        """
        When two detections overlap (e.g., "John" as PERSON and part
        of an EMAIL_ADDRESS), keep only the one with the higher score.
        If scores are equal, keep the longer span.
        """
        if not results:
            return []

        # Sort by start position, then by score descending
        sorted_results = sorted(
            results, key=lambda r: (r.start, -r.score)
        )
        merged: List[RecognizerResult] = [sorted_results[0]]

        for current in sorted_results[1:]:
            previous = merged[-1]
            # Check for overlap
            if current.start < previous.end:
                # Keep the one with higher score; on tie, keep longer span
                if current.score > previous.score or (
                    current.score == previous.score
                    and (current.end - current.start)
                    > (previous.end - previous.start)
                ):
                    merged[-1] = current
                # else: keep previous (already in merged)
            else:
                merged.append(current)

        return merged


# ──────────────────────────────────────────────────────────────────────
# 4. DOCX Processor  —  format-preserving document walker
# ──────────────────────────────────────────────────────────────────────

class DocxProcessor:
    """
    Walks a .docx document and applies PII redaction while strictly
    preserving the original formatting.

    Strategy:
        - python-docx represents styled text as *runs* within a
          paragraph.  Each run has its own font properties.
        - We concatenate run texts to form the full paragraph, run
          the redaction engine on the concatenated string, then
          redistribute the redacted characters back into the original
          run boundaries.
        - This ensures bold / italic / colour / font-size are NOT
          disturbed.
    """

    def __init__(self, engine: RedactionEngine) -> None:
        self.engine = engine
        self._stats = {"paragraphs": 0, "tables": 0, "detections": 0}

    def process(self, input_path: str, output_path: str) -> Dict:
        """
        Read *input_path*, redact PII, and write to *output_path*.
        Returns processing statistics.
        """
        doc = Document(input_path)

        # --- Body paragraphs ----------------------------------------------
        for paragraph in doc.paragraphs:
            self._redact_paragraph(paragraph)

        # --- Floating Text Boxes ------------------------------------------
        # Deep XML traversal to find w:txbxContent nodes
        for txbx in doc.element.xpath('.//w:txbxContent'):
            for p_elm in txbx.xpath('.//w:p'):
                para = Paragraph(p_elm, doc)
                self._redact_paragraph(para)

        # --- OCR Embedded Images ------------------------------------------
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_bytes))
                    extracted_text = pytesseract.image_to_string(img)
                    if extracted_text.strip():
                        res = self.engine.redact(extracted_text)
                        if res.detections:
                            print(f"📸 [OCR] Found PII in image {rel.target_ref}:")
                            for d in res.detections:
                                val = extracted_text[d.start:d.end].strip()
                                print(f"    - {d.entity_type}: {val!r}")
                except Exception as e:
                    print(f"⚠️ [OCR] Error processing image {rel.target_ref}: {e}")

        # --- Headers & Footers --------------------------------------------
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                self._redact_paragraph(header_para)
            for footer_para in section.footer.paragraphs:
                self._redact_paragraph(footer_para)
            # Also handle tables inside headers / footers
            for tbl in section.header.tables:
                self._redact_table(tbl)
            for tbl in section.footer.tables:
                self._redact_table(tbl)

        # --- Tables -------------------------------------------------------
        for table in doc.tables:
            self._redact_table(table)

        doc.save(output_path)

        return {
            "input": input_path,
            "output": output_path,
            "paragraphs_processed": self._stats["paragraphs"],
            "tables_processed": self._stats["tables"],
            "total_detections": self._stats["detections"],
            "mapping": self.engine.mapper.get_full_map(),
        }

    # ---- internal helpers ------------------------------------------------

    def _redact_paragraph(self, paragraph) -> None:
        """
        Redact PII in a single paragraph while preserving run-level
        formatting.

        Algorithm:
            1. Collect all runs and concatenate their .text to get the
               full paragraph text.
            2. Run the RedactionEngine on the full text.
            3. Re-distribute the redacted text back across the original
               run boundaries.  If the redacted text is shorter or
               longer than the original, the extra characters are
               absorbed by / appended to the last affected run.
        """
        runs = paragraph.runs
        if not runs:
            return

        # Build full text and a mapping of (start, end) per run
        full_text = ""
        run_boundaries: List[Tuple[int, int]] = []
        for run in runs:
            start = len(full_text)
            full_text += run.text
            run_boundaries.append((start, len(full_text)))

        if not full_text.strip():
            return

        # Redact the concatenated text
        result = self.engine.redact(full_text)
        self._stats["paragraphs"] += 1
        self._stats["detections"] += len(result.detections)

        if result.redacted_text == full_text:
            return  # Nothing changed — skip write to avoid any risk

        # Redistribute redacted text into existing runs
        redacted = result.redacted_text
        self._redistribute_text_to_runs(runs, run_boundaries, full_text, redacted)

    def _redistribute_text_to_runs(
        self,
        runs,
        run_boundaries: List[Tuple[int, int]],
        original_full: str,
        redacted_full: str,
    ) -> None:
        """
        Map the redacted full text back onto the original runs.

        We identify which character ranges changed (due to entity
        replacements that may differ in length) and carefully assign
        characters to each run so formatting stays intact.
        """
        # Build a character-level mapping from original positions to
        # redacted positions using the detection spans.
        # Simpler approach: walk detections, build offset-adjusted
        # redacted string segments for each run.

        # For robustness we use a simple proportional approach:
        #   - If no length change occurred, slice directly.
        #   - If length changed, we attempt smart redistribution.

        if len(redacted_full) == len(original_full):
            # Same length — easy 1:1 character mapping
            for i, run in enumerate(runs):
                start, end = run_boundaries[i]
                run.text = redacted_full[start:end]
            return

        # Length changed — we need to figure out which runs are affected.
        # Strategy: Find replacement regions, adjust run texts accordingly.
        # We rebuild run texts by tracking cumulative offset shifts.

        detections = sorted(
            self.engine.redact(original_full).detections,
            key=lambda d: d.start,
        )

        # Build a list of (orig_start, orig_end, replacement_text)
        replacements: List[Tuple[int, int, str]] = []
        for det in detections:
            orig_val = original_full[det.start : det.end]
            repl_val = self.engine.mapper.get_replacement(
                det.entity_type, orig_val
            )
            replacements.append((det.start, det.end, repl_val))

        # Now rebuild each run's text accounting for replacements
        offset = 0  # cumulative shift
        repl_idx = 0
        for i, run in enumerate(runs):
            r_start, r_end = run_boundaries[i]
            run_text = run.text
            new_run_text = ""
            pos = r_start  # position in original full text

            while pos < r_end:
                # Check if a replacement starts at or before current pos
                if repl_idx < len(replacements):
                    rep_start, rep_end, rep_text = replacements[repl_idx]
                    if rep_start >= r_start and rep_start < r_end:
                        # Replacement starts in this run
                        # Add any text before the replacement
                        new_run_text += original_full[pos:rep_start]
                        if rep_end <= r_end:
                            # Replacement fits entirely in this run
                            new_run_text += rep_text
                            pos = rep_end
                            repl_idx += 1
                            continue
                        else:
                            # Replacement spans into next run(s)
                            # Put full replacement in this run
                            new_run_text += rep_text
                            pos = r_end
                            # Advance repl_idx only when we've consumed
                            # all original chars of this replacement
                            # (handled in subsequent runs)
                            # Mark how many chars remain
                            replacements[repl_idx] = (
                                r_end,
                                rep_end,
                                "",  # already emitted the text
                            )
                            continue
                    elif rep_start < r_start and rep_end > r_start:
                        # Replacement started in previous run, spans here
                        consumed = min(rep_end, r_end)
                        # Text was already emitted; skip original chars
                        pos = consumed
                        if consumed >= rep_end:
                            repl_idx += 1
                        else:
                            replacements[repl_idx] = (
                                consumed, rep_end, ""
                            )
                        continue

                # No replacement at this position — copy original char
                new_run_text += original_full[pos]
                pos += 1

            run.text = new_run_text

    def _redact_table(self, table) -> None:
        """Iterate every cell in every row of a table and redact."""
        self._stats["tables"] += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._redact_paragraph(paragraph)
                # Handle nested tables (yes, .docx supports them)
                for nested_table in cell.tables:
                    self._redact_table(nested_table)


# ──────────────────────────────────────────────────────────────────────
# 5. Evaluation Engine  —  Precision / Recall / Accuracy
# ──────────────────────────────────────────────────────────────────────

@dataclass
class Annotation:
    """A single ground-truth PII annotation."""
    start: int
    end: int
    entity_type: str
    text: str


@dataclass
class EvaluationMetrics:
    """Computed evaluation metrics."""
    true_positives: int = 0
    false_positives: int = 0
    false_negatives: int = 0
    true_negatives: int = 0  # meaningful only at token level

    @property
    def precision(self) -> float:
        """Of everything the model flagged, what fraction was correct?"""
        denom = self.true_positives + self.false_positives
        return self.true_positives / denom if denom > 0 else 0.0

    @property
    def recall(self) -> float:
        """Of all real PII, what fraction did the model catch?"""
        denom = self.true_positives + self.false_negatives
        return self.true_positives / denom if denom > 0 else 0.0

    @property
    def f1(self) -> float:
        """Harmonic mean of precision and recall."""
        p, r = self.precision, self.recall
        return 2 * p * r / (p + r) if (p + r) > 0 else 0.0

    @property
    def accuracy(self) -> float:
        """
        Overall accuracy across all classification decisions.
        (TP + TN) / (TP + TN + FP + FN)
        """
        total = (
            self.true_positives
            + self.true_negatives
            + self.false_positives
            + self.false_negatives
        )
        return (
            (self.true_positives + self.true_negatives) / total
            if total > 0
            else 0.0
        )

    def summary(self) -> str:
        """Human-readable metrics summary with Confusion Matrix."""
        return (
            f"╔══════════════════════════════════════╗\n"
            f"║       EVALUATION METRICS REPORT      ║\n"
            f"╠══════════════════════════════════════╣\n"
            f"║  True Positives  : {self.true_positives:>6}            ║\n"
            f"║  False Positives : {self.false_positives:>6}            ║\n"
            f"║  False Negatives : {self.false_negatives:>6}            ║\n"
            f"║  True Negatives  : {self.true_negatives:>6}            ║\n"
            f"╠══════════════════════════════════════╣\n"
            f"║  Precision       : {self.precision:>6.2%}            ║\n"
            f"║  Recall          : {self.recall:>6.2%}            ║\n"
            f"║  F1 Score        : {self.f1:>6.2%}            ║\n"
            f"║  Accuracy        : {self.accuracy:>6.2%}            ║\n"
            f"╚══════════════════════════════════════╝\n"
            f"\n"
            f"--- CONFUSION MATRIX ---\n"
            f"                 Predicted PII | Predicted Non-PII\n"
            f"Actual PII     | {self.true_positives:>13} | {self.false_negatives:>17}\n"
            f"Actual Non-PII | {self.false_positives:>13} | {self.true_negatives:>17}\n"
        )


class EvaluationEngine:
    """
    Compares Presidio detections against a ground-truth annotation set
    and computes Precision, Recall, F1, and Accuracy.

    Ground-Truth Format (JSON):
        {
            "annotations": [
                {
                    "start": 0,
                    "end": 12,
                    "entity_type": "PERSON",
                    "text": "Rashi Patil"
                },
                ...
            ],
            "full_text": "Rashi Patil works at ..."
        }

    Matching Strategy:
        A detection is a True Positive if it overlaps ≥50% of a
        ground-truth span AND has the same entity type.  This
        tolerates minor boundary differences (e.g., capturing a
        trailing space).
    """

    OVERLAP_THRESHOLD: float = 0.50  # minimum IoU for a match

    def __init__(self, ground_truth_path: str) -> None:
        with open(ground_truth_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        self.full_text: str = data["full_text"]
        self.annotations: List[Annotation] = [
            Annotation(**a) for a in data["annotations"]
        ]

    def evaluate(
        self, detections: List[RecognizerResult]
    ) -> EvaluationMetrics:
        """
        Compare *detections* against ground-truth and return metrics.
        """
        metrics = EvaluationMetrics()
        matched_annotations: Set[int] = set()
        matched_detections: Set[int] = set()

        # --- Match detections to annotations -----------------------------
        for d_idx, detection in enumerate(detections):
            best_match: Optional[int] = None
            best_overlap: float = 0.0

            for a_idx, annotation in enumerate(self.annotations):
                if a_idx in matched_annotations:
                    continue
                if detection.entity_type != annotation.entity_type:
                    continue

                overlap = self._compute_overlap(
                    detection.start, detection.end,
                    annotation.start, annotation.end,
                )
                if overlap >= self.OVERLAP_THRESHOLD and overlap > best_overlap:
                    best_overlap = overlap
                    best_match = a_idx

            if best_match is not None:
                metrics.true_positives += 1
                matched_annotations.add(best_match)
                matched_detections.add(d_idx)
            else:
                metrics.false_positives += 1

        # --- Count missed annotations (false negatives) ------------------
        metrics.false_negatives = len(self.annotations) - len(
            matched_annotations
        )

        # --- Estimate true negatives (token-level) -----------------------
        # We approximate TN by counting whitespace-delimited tokens that
        # are neither in a ground-truth span nor in a detection span.
        all_chars = set(range(len(self.full_text)))
        gt_chars = set()
        for a in self.annotations:
            gt_chars.update(range(a.start, a.end))
        det_chars = set()
        for d in detections:
            det_chars.update(range(d.start, d.end))
        non_pii_chars = all_chars - gt_chars
        correctly_ignored = non_pii_chars - det_chars
        # Convert to approximate token count
        non_pii_text = "".join(
            self.full_text[i] if i in correctly_ignored else " "
            for i in range(len(self.full_text))
        )
        metrics.true_negatives = len(non_pii_text.split())

        return metrics

    @staticmethod
    def _compute_overlap(
        s1: int, e1: int, s2: int, e2: int
    ) -> float:
        """
        Compute Intersection-over-Union (IoU) for two character spans.
        """
        intersection = max(0, min(e1, e2) - max(s1, s2))
        union = max(e1, e2) - min(s1, s2)
        return intersection / union if union > 0 else 0.0

    @classmethod
    def evaluate_from_text(
        cls,
        text: str,
        ground_truth_path: str,
        analyzer: AnalyzerEngine,
    ) -> EvaluationMetrics:
        """
        Convenience method: analyse *text*, then evaluate against
        ground-truth in a single call.
        """
        detections = analyzer.analyze(
            text=text,
            entities=SUPPORTED_PII_TYPES,
            language="en",
        )
        engine = cls(ground_truth_path)
        return engine.evaluate(detections)

    def update_report(self, metrics: EvaluationMetrics, report_path: str = "EVALUATION_REPORT.md") -> None:
        """Update EVALUATION_REPORT.md with the latest Confusion Matrix and metrics."""
        try:
            with open(report_path, "r", encoding="utf-8") as f:
                content = f.read()

            summary_str = metrics.summary()
            
            new_content = re.sub(
                r"(### 3\.1 Aggregate Metrics\n\n).*?(?=\n\n> \*\*Note\*\*:)",
                r"\1```\n" + summary_str + r"```",
                content,
                flags=re.DOTALL
            )
            
            with open(report_path, "w", encoding="utf-8") as f:
                f.write(new_content)
            print(f"✅ Successfully updated {report_path} with latest metrics and Confusion Matrix.")
        except Exception as e:
            print(f"⚠️ Failed to update {report_path}: {e}")


# ──────────────────────────────────────────────────────────────────────
# 6. CLI Entry Point
# ──────────────────────────────────────────────────────────────────────

def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Enterprise PII Redaction Engine for .docx files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        "--input", "-i",
        required=True,
        help="Path to the input .docx file.",
    )
    parser.add_argument(
        "--output", "-o",
        required=True,
        help="Path to write the redacted .docx file.",
    )
    parser.add_argument(
        "--ground-truth", "-g",
        default=None,
        help="Path to a ground-truth annotations JSON file for evaluation.",
    )
    parser.add_argument(
        "--evaluate", "-e",
        action="store_true",
        help="Run the evaluation engine (requires --ground-truth).",
    )
    parser.add_argument(
        "--threshold", "-t",
        type=float,
        default=CONFIDENCE_THRESHOLD,
        help=f"Confidence threshold (default: {CONFIDENCE_THRESHOLD}).",
    )
    parser.add_argument(
        "--mapping-output", "-m",
        default=None,
        help="Path to save the 1:1 PII mapping dictionary as JSON.",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()

    # Validate inputs
    if not Path(args.input).exists():
        print(f"❌ Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)
    if args.evaluate and not args.ground_truth:
        print(
            "❌ --evaluate requires --ground-truth <path>",
            file=sys.stderr,
        )
        sys.exit(1)

    print("🔧 Initialising Presidio + spaCy engine …")
    analyzer = _build_analyzer()
    mapper = ConsistentMapper()
    engine = RedactionEngine(
        analyzer=analyzer,
        mapper=mapper,
        confidence_threshold=args.threshold,
    )
    processor = DocxProcessor(engine)

    print(f"📄 Processing: {args.input}")
    stats = processor.process(args.input, args.output)

    print(f"✅ Redacted document saved to: {args.output}")
    print(f"   Paragraphs processed : {stats['paragraphs_processed']}")
    print(f"   Tables processed     : {stats['tables_processed']}")
    print(f"   Total PII detections : {stats['total_detections']}")

    # Save mapping dictionary if requested
    if args.mapping_output:
        with open(args.mapping_output, "w", encoding="utf-8") as f:
            json.dump(stats["mapping"], f, indent=2)
        print(f"📋 Mapping dictionary saved to: {args.mapping_output}")

    # Print mapping summary
    print("\n📋 PII Replacement Mapping:")
    print("─" * 60)
    for entity_type, mappings in stats["mapping"].items():
        print(f"  [{entity_type}]")
        for original, replacement in mappings.items():
            print(f"    {original!r:40s} → {replacement!r}")
    print("─" * 60)

    # Run evaluation if requested
    if args.evaluate:
        print("\n📊 Running evaluation against ground truth …")
        eval_engine = EvaluationEngine(args.ground_truth)
        # Re-analyse the full text to get detections for evaluation
        doc = Document(args.input)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        detections = analyzer.analyze(
            text=full_text,
            entities=SUPPORTED_PII_TYPES,
            language="en",
        )
        metrics = eval_engine.evaluate(detections)
        print(metrics.summary())
        # Automatically update the evaluation report
        eval_engine.update_report(metrics)


if __name__ == "__main__":
    main()
