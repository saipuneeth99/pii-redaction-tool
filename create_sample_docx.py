"""
create_sample_docx.py — Generate a sample .docx file loaded with PII
for testing the redaction engine.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Title ──
title = doc.add_heading("CONFIDENTIAL — Employee Onboarding Record", level=1)

# ── Section 1: Personal Information ──
doc.add_heading("1. Personal Information", level=2)

p1 = doc.add_paragraph()
run = p1.add_run("Full Name: ")
run.bold = True
p1.add_run("Rashi Patil")

p2 = doc.add_paragraph()
run = p2.add_run("Date of Birth: ")
run.bold = True
p2.add_run("March 15, 1990")

p3 = doc.add_paragraph()
run = p3.add_run("Social Security Number: ")
run.bold = True
p3.add_run("123-45-6789")

p4 = doc.add_paragraph()
run = p4.add_run("Email Address: ")
run.bold = True
run2 = p4.add_run("rashi.patil@acmecorp.com")
run2.font.color.rgb = RGBColor(0, 0, 255)

p5 = doc.add_paragraph()
run = p5.add_run("Phone: ")
run.bold = True
p5.add_run("(555) 123-4567")

p6 = doc.add_paragraph()
run = p6.add_run("Home Address: ")
run.bold = True
p6.add_run("742 Evergreen Terrace, Springfield, IL 62704")

# ── Section 2: Employment Details ──
doc.add_heading("2. Employment Details", level=2)

p7 = doc.add_paragraph()
run = p7.add_run("Employer: ")
run.bold = True
run2 = p7.add_run("Acme Corp")
run2.italic = True

p8 = doc.add_paragraph()
p8.add_run("Rashi Patil has been assigned to the Engineering division. "
           "Her corporate credit card number is ")
run_cc = p8.add_run("4532015112830366")
run_cc.bold = True
run_cc.font.color.rgb = RGBColor(255, 0, 0)
p8.add_run(" for business travel expenses.")

p9 = doc.add_paragraph()
p9.add_run("Her workstation IP address is ")
run_ip = p9.add_run("192.168.1.100")
run_ip.font.name = "Courier New"
run_ip.font.size = Pt(10)
p9.add_run(". IT support can be contacted at support@acmecorp.com.")

# ── Section 3: Emergency Contact ──
doc.add_heading("3. Emergency Contact", level=2)

# Table with formatting
table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
table.cell(0, 0).text = "Field"
table.cell(0, 1).text = "Details"
table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

table.cell(1, 0).text = "Contact Name"
table.cell(1, 1).text = "Vikram Sharma"

table.cell(2, 0).text = "Phone"
table.cell(2, 1).text = "(555) 987-6543"

table.cell(3, 0).text = "Email"
table.cell(3, 1).text = "vikram.sharma@gmail.com"

# ── Section 4: Additional Notes ──
doc.add_heading("4. Notes", level=2)
p10 = doc.add_paragraph()
p10.add_run("Please ensure that Rashi Patil's records are stored securely. "
            "A copy of this document was sent to legal@acmecorp.com and "
            "Vikram Sharma on 08/15/2024. "
            "The backup server is accessible at 10.0.0.42.")

doc.save("sample_input.docx")
print("✅ sample_input.docx created successfully.")
